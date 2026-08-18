#!/usr/bin/env python3
"""Sinh file .docx từ một tài liệu Markdown trong `docs/`.

Vì sao có script này thay vì gõ tay một lần: tài liệu hướng dẫn sẽ còn sửa
nhiều lượt, và mỗi lượt lại phải xuất bản Word cho người không đọc Markdown.
Sửa .md rồi chạy lại script là xong — không có chuyện hai bản lệch nhau vì ai
đó sửa .md mà quên .docx.

Bench này KHÔNG có pandoc/libreoffice (đã kiểm 18/08/2026); chỉ có python-docx.

    python3 docs/md2docx.py docs/HDSD-phan-quyen-khoa-phong.md
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

DAM = re.compile(r"\*\*(.+?)\*\*")
MA = re.compile(r"`([^`]+)`")


def _viet_doan(p, text, dam=False, nghieng=False):
    """Đổ text vào một paragraph, giữ **đậm**, *nghiêng* và `mã`.

    ĐỆ QUY, không phải một lượt tách phẳng: định dạng lồng nhau là chuyện
    thường trong tài liệu này (`**`Portal Member`**`, `**`CHUNG` là mã dành
    riêng**`). Bản trước tách một lượt nên khối đậm được đổ nguyên văn kèm
    dấu backtick lòi ra bản Word — lỗi bắt được khi rà file .docx sinh ra,
    không phải khi đọc mã.
    """
    for phan in re.split(r"(\*\*.+?\*\*|\*[^*\n]+?\*|`[^`]+`)", text):
        if not phan:
            continue
        if phan.startswith("**") and phan.endswith("**") and len(phan) > 4:
            _viet_doan(p, phan[2:-2], dam=True, nghieng=nghieng)
        elif phan.startswith("*") and phan.endswith("*") and len(phan) > 2:
            _viet_doan(p, phan[1:-1], dam=dam, nghieng=True)
        elif phan.startswith("`") and phan.endswith("`") and len(phan) > 2:
            r = p.add_run(phan[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xB0, 0x30, 0x30)
            r.bold, r.italic = dam, nghieng
        else:
            r = p.add_run(phan)
            r.bold, r.italic = dam, nghieng


def _bang(doc, dong):
    """`dong` là danh sách dòng markdown của MỘT bảng, kể cả dòng gạch ngang."""
    hang = [[o.strip() for o in d.strip().strip("|").split("|")] for d in dong]
    hang = [h for h in hang if not all(set(o) <= set("-: ") for o in h)]
    if not hang:
        return
    so_cot = max(len(h) for h in hang)
    t = doc.add_table(rows=0, cols=so_cot)
    t.style = "Table Grid"
    for i, h in enumerate(hang):
        o = t.add_row().cells
        for j in range(so_cot):
            o[j].text = ""
            _viet_doan(o[j].paragraphs[0], h[j] if j < len(h) else "")
            if i == 0:
                for r in o[j].paragraphs[0].runs:
                    r.bold = True


def chuyen(md_path: Path) -> Path:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    dong = md_path.read_text(encoding="utf-8").split("\n")
    i = 0
    while i < len(dong):
        d = dong[i]

        if d.startswith("```"):  # khối mã
            i += 1
            khoi = []
            while i < len(dong) and not dong[i].startswith("```"):
                khoi.append(dong[i])
                i += 1
            p = doc.add_paragraph()
            r = p.add_run("\n".join(khoi))
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            p.paragraph_format.left_indent = Pt(18)
            i += 1
            continue

        if d.startswith("|"):  # bảng
            khoi = []
            while i < len(dong) and dong[i].startswith("|"):
                khoi.append(dong[i])
                i += 1
            _bang(doc, khoi)
            doc.add_paragraph()
            continue

        if d.startswith("#"):
            cap = len(d) - len(d.lstrip("#"))
            doc.add_heading(d.lstrip("#").strip(), level=min(cap, 4))
        elif d.strip() == "---":
            doc.add_paragraph("─" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif d.startswith("> "):
            khoi = []
            while i < len(dong) and dong[i].startswith(">"):
                khoi.append(dong[i].lstrip(">").strip())
                i += 1
            p = doc.add_paragraph(style="Intense Quote")
            _viet_doan(p, " ".join(x for x in khoi if x))
            continue
        elif re.match(r"^\s*(\d+\.|-) ", d):
            # Một mục danh sách có thể trải nhiều dòng nguồn. Không gom lại
            # thì một cụm `**đậm**` bắc qua hai dòng sẽ lòi dấu sao ra bản
            # Word — đúng lỗi đã bắt được ở lượt sinh đầu tiên.
            danh_so = bool(re.match(r"^\s*\d+\. ", d))
            khoi = [re.sub(r"^\s*(\d+\.|-) ", "", d).strip()]
            j = i + 1
            while j < len(dong) and dong[j].strip() and not re.match(
                r"^(#|\||```|> |\s*-\s|\s*\d+\. |---$)", dong[j]
            ):
                khoi.append(dong[j].strip())
                j += 1
            p = doc.add_paragraph(style="List Number" if danh_so else "List Bullet")
            _viet_doan(p, " ".join(khoi))
            i = j
            continue
        elif d.strip():
            # Gom các dòng văn xuôi liên tiếp thành MỘT đoạn. Markdown xuống
            # dòng trong nguồn chỉ là để giữ độ rộng cột chữ, không phải ngắt
            # đoạn — bản trước đổ mỗi dòng thành một paragraph nên văn xuôi
            # trong Word bị vụn ra từng mẩu ba bốn chữ.
            khoi = [d.strip()]
            j = i + 1
            while j < len(dong) and dong[j].strip() and not re.match(
                r"^(#|\||```|> |- |\d+\. |---$)", dong[j]
            ):
                khoi.append(dong[j].strip())
                j += 1
            p = doc.add_paragraph()
            _viet_doan(p, " ".join(khoi))
            i = j
            continue
        i += 1

    out = md_path.with_suffix(".docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    if len(sys.argv) != 2:
        sys.exit("Dùng: python3 docs/md2docx.py <đường-dẫn-file.md>")
    print("Đã sinh:", chuyen(Path(sys.argv[1])))
